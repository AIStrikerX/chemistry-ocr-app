import streamlit as st
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import cv2
import numpy as np
import os
import ssl
try:
    import streamlit.secret as secrets
except ImportError:
    pass

# WORKAROUND: Bypass SSL verification for model downloads if certificates are missing
try:
    _create_unverified_https_context = ssl._create_unverified_context
except AttributeError:
    pass
else:
    ssl._create_default_https_context = _create_unverified_https_context

# Set page configuration
st.set_page_config(page_title="Smart OCR - Chemistry Notes", page_icon="🧪", layout="wide")

st.markdown("""
<style>
    .main { background-color: #f0f2f6; }
    .stButton>button { color: white; background-color: #4CAF50; border-radius: 5px; }
</style>
""", unsafe_allow_html=True)

st.title("🧪 Smart OCR - Chemistry Notes")
st.markdown("**Layout-aware OCR** for handwritten chemistry notes with structured Word export")

# Sidebar for engine selection
st.sidebar.title("⚙️ OCR Settings")
ocr_engine = st.sidebar.radio(
    "Select OCR Engine:",
    ["Groq (Llama 4 Maverick)", "Gemini AI (Best Quality)"],
    help="Groq Llama 4 is fast and free (preview). Gemini requires API key."
)

# API Keys
gemini_api_key = None
groq_api_key = None

if ocr_engine == "Gemini AI (Best Quality)":
    gemini_api_key = st.sidebar.text_input(
        "Gemini API Key:",
        type="password",
        help="Get your free key at: https://aistudio.google.com/app/apikey"
    )
    if not gemini_api_key:
        st.sidebar.warning("⚠️ API key required for Gemini")

elif ocr_engine == "Groq (Llama 4 Maverick)":
    # Try to load from secrets first
    if "GROQ_API_KEY" in st.secrets:
        groq_api_key = st.secrets["GROQ_API_KEY"]
    else:
        groq_api_key = st.sidebar.text_input(
            "Groq API Key:",
            type="password",
            help="Get your free key at: https://console.groq.com/keys"
        )
    
    transcription_mode = st.sidebar.radio(
        "Transcription Mode:",
        ["Relaxed (Clean Notes)", "Strict (Exact Transcription)"],
        help="Relaxed fixes minor errors. Strict keeps text exactly as written."
    )
    
    if not groq_api_key:
        st.sidebar.warning("⚠️ API key required for Groq")


# Initialize engines
@st.cache_resource
def load_groq_client(api_key):
    from groq import Groq
    return Groq(api_key=api_key)

@st.cache_resource
def load_gemini(_api_key):
    import google.generativeai as genai
    genai.configure(api_key=_api_key)
    return genai.GenerativeModel('gemini-1.5-flash')

def encode_image(image_path):
    import base64
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

def groq_ocr_process(image_path, api_key, mode="Relaxed (Clean Notes)"):
    """Process image with Groq Llama 4 Maverick using improved system prompt"""
    try:
        client = load_groq_client(api_key)
        base64_image = encode_image(image_path)
        
        # Base System Prompt
        system_prompt = """You are an expert scientific document transcriber and editor.

Your task is to convert handwritten chemistry notes from an image into clean, structured study notes.

IMPORTANT CONTEXT:
- The source is a photographed handwritten notebook page.
- The page may contain multiple columns, boxed sections, underlines, arrows, diagrams, equations, and headings written in different colors.
- Some content is textual, some is visual (diagrams, flowcharts, graphs).

YOUR RESPONSIBILITIES:
1. Accurately transcribe ALL readable text from the image.
2. Preserve the original meaning, terminology, and scientific correctness.
3. Reconstruct a logical structure similar to well-written chemistry notes.

STRUCTURE RULES:
- Use clear section headings for major topics.
- Use subheadings where appropriate.
- Use bullet points or numbered lists when the content implies lists.
- Maintain logical reading order (top-to-bottom, left-to-right).
- Do NOT invent new content or explanations.

EQUATIONS & SYMBOLS:
- Preserve chemical symbols, formulas, charges, arrows, and proportionality signs.
- Use LaTeX-style inline math only when necessary (e.g., H₂SO₄, V₂O₅, 1/viscosity).
- Do NOT “correct” chemistry unless the handwriting is clearly ambiguous.

DIAGRAMS & VISUAL ELEMENTS:
- If a diagram, graph, flowchart, or visual illustration is present:
  - Do NOT attempt to recreate it as text.
  - Insert a placeholder in the format:
    [DIAGRAM: short factual description of what is shown]
- If arrows indicate process flow, reflect that flow in text where obvious.

QUALITY CONTROL:
- If a word is unclear, transcribe the closest plausible chemistry term without guessing new concepts.
- Do not paraphrase or summarize aggressively.
- Do not beautify language; keep it note-like and concise.

OUTPUT FORMAT:
- Output clean, structured Markdown.
- Do NOT include explanations about what you are doing.
- Do NOT mention OCR, AI, or the model.
- Output ONLY the final structured notes."""

        # Add Mode Logic
        if "Strict" in mode:
            system_prompt += "\n\nSTRICT MODE: If any text is unclear, preserve it as written rather than guessing."
        else:
            system_prompt += "\n\nRELAXED MODE: Minor spelling corrections are allowed only for standard chemistry terms."

        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": system_prompt
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Transcribe and structure the handwritten chemistry notes in this image according to the system instructions."},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                            },
                        },
                    ],
                }
            ],
            model="meta-llama/llama-4-maverick-17b-128e-instruct",
            temperature=0.2, # Added reliability constraint
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        st.error(f"Groq Error: {e}")
        return None

def preprocess_image(image_path):
    """Enhanced preprocessing for better OCR"""
    img = cv2.imread(image_path)
    if img is None:
        return None
    
    # Convert to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Remove noise
    denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
    
    # Adaptive threshold
    thresh = cv2.adaptiveThreshold(
        denoised, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 11, 2
    )
    
    return thresh

def gemini_ocr_process(image_path, api_key):
    """Process image with Gemini Vision API"""
    try:
        model = load_gemini(api_key)
        img = Image.open(image_path)
        
        prompt = """You are an expert transcriber for handwritten chemistry notes.

Please transcribe this handwritten chemistry note into clean, structured Markdown format.

Follow these rules:
1. Preserve the document structure (headings, paragraphs, lists)
2. Use proper Markdown headings (# for main topics, ## for subtopics)
3. Format chemical equations properly
4. If you see diagrams or complex equations, describe them as: `[DIAGRAM: brief description]`
5. Maintain the reading order (left-to-right, top-to-bottom)
6. Fix any obvious spelling errors in chemistry terms

Output ONLY the Markdown text, no additional commentary."""

        response = model.generate_content([prompt, img])
        return response.text
    except Exception as e:
        st.error(f"Gemini Error: {e}")
        return None

def create_structured_docx(ocr_result, image_path, engine_type="groq"):
    """Create a structured Word document from OCR results"""
    doc = Document()
    
    # Title
    title = doc.add_heading("Chemistry Notes - OCR Extraction", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"OCR Engine: {engine_type.upper()}")
    doc.add_paragraph("_" * 50)
    
    if engine_type in ["groq", "gemini"]:
        # Markdown output - convert to Word
        if ocr_result:
            lines = ocr_result.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                    
                # Detect markdown headings
                if line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('[DIAGRAM'):
                    doc.add_paragraph(line, style='Intense Quote')
                else:
                    doc.add_paragraph(line)
    
    # Add separator
    doc.add_paragraph()
    doc.add_paragraph("_" * 50)
    
    # Add original image at the end for reference
    doc.add_heading("Original Image (Reference)", level=2)
    try:
        doc.add_picture(image_path, width=Inches(6))
    except:
        doc.add_paragraph("[Could not embed original image]")
    
    return doc


# Main app
uploaded_file = st.file_uploader("📤 Upload handwritten note", type=["jpg", "jpeg", "png"])

if uploaded_file:
    col1, col2 = st.columns([1, 1])
    
    # Save uploaded file
    with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
        tmp.write(uploaded_file.read())
        input_path = tmp.name
    
    with col1:
        st.subheader("📸 Original Image")
        st.image(input_path, use_container_width=True)
    
    with col2:
        st.subheader("📝 Extracted Content")
        
        if ocr_engine == "Groq (Llama 4 Maverick)":
            if groq_api_key:
                with st.spinner('⚡ Groq (Llama 4 Maverick) is analyzing your note...'):
                    ocr_result = groq_ocr_process(input_path, groq_api_key, transcription_mode)
                    
                    if ocr_result:
                        st.markdown(ocr_result)
                        doc = create_structured_docx(ocr_result, input_path, engine_type="groq")
                    else:
                        st.error("Failed to process with Groq. Check API Key or try again.")
                        doc = None
            else:
                st.warning("⚠️ Please enter your Groq API key in the sidebar")
                doc = None
            
        else:  # Gemini
            if gemini_api_key:
                with st.spinner('🤖 Gemini AI is analyzing your note...'):
                    ocr_result = gemini_ocr_process(input_path, gemini_api_key)
                    
                    if ocr_result:
                        st.markdown(ocr_result)
                        doc = create_structured_docx(ocr_result, input_path, engine_type="gemini")
                    else:
                        st.error("Failed to process with Gemini")
                        doc = None
            else:
                st.warning("⚠️ Please enter your Gemini API key in the sidebar")
                doc = None
    
    # Download button
    if 'doc' in locals() and doc is not None:
        doc_path = "chemistry_notes_ocr.docx"
        doc.save(doc_path)
        
        with open(doc_path, "rb") as f:
            st.download_button(
                label="⬇️ Download Structured Word Document",
                data=f,
                file_name="chemistry_notes_structured.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        # Cleanup
        try:
            os.remove(input_path)
        except:
            pass

st.sidebar.markdown("---")
st.sidebar.markdown("""
### 📚 Tips for Best Results
- **Lighting**: Ensure good, even lighting
- **Focus**: Clear, sharp images work best
- **Resolution**: Higher resolution = better accuracy
- **Groq Llama 4**: Extremely fast, multimodal
- **Gemini**: Excellent for complex layouts

### 🧪 Chemistry Notes
- Diagrams will be marked as `[DIAGRAM]`
- Equations are formatted in Markdown
- Original image embedded for reference
""")
